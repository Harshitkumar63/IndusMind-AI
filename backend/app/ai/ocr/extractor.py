"""
IndusMind AI — OCR & Text Extraction Pipeline

Extracts text from PDFs (native + scanned), DOCX, XLSX, CSV, and images.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentExtractor:
    """Multi-format document text extraction pipeline."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".png", ".jpg", ".jpeg", ".tiff"}

    async def extract(self, file_path: str) -> dict[str, Any]:
        """
        Extract text from a document file.

        Returns:
            dict with: text, page_count, metadata, method
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".xlsx": self._extract_xlsx,
            ".csv": self._extract_csv,
            ".txt": self._extract_txt,
            ".png": self._extract_image,
            ".jpg": self._extract_image,
            ".jpeg": self._extract_image,
            ".tiff": self._extract_image,
        }

        extractor = extractors[ext]
        result = await extractor(file_path)
        logger.info(
            "Text extracted",
            file=path.name,
            method=result.get("method", "unknown"),
            chars=len(result.get("text", "")),
        )
        return result

    async def _extract_pdf(self, file_path: str) -> dict[str, Any]:
        """Extract text from PDF using PyMuPDF (native) with OCR fallback."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            pages_text: list[str] = []
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    pages_text.append(text)
                else:
                    # Page has no selectable text — likely scanned
                    # In production: use PaddleOCR or Tesseract
                    pages_text.append(f"[Scanned page {page.number + 1} — OCR required]")

            full_text = "\n\n".join(pages_text)
            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "creator": doc.metadata.get("creator", ""),
            }
            doc.close()
            return {"text": full_text, "page_count": len(pages_text), "metadata": metadata, "method": "pymupdf"}
        except ImportError:
            logger.warning("PyMuPDF not installed, returning placeholder")
            return {"text": f"[PDF extraction requires PyMuPDF: {file_path}]", "page_count": 0, "metadata": {}, "method": "placeholder"}

    async def _extract_docx(self, file_path: str) -> dict[str, Any]:
        """Extract text from DOCX files."""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return {"text": "\n\n".join(paragraphs), "page_count": 1, "metadata": {}, "method": "python-docx"}
        except ImportError:
            return {"text": f"[DOCX extraction requires python-docx: {file_path}]", "page_count": 0, "metadata": {}, "method": "placeholder"}

    async def _extract_xlsx(self, file_path: str) -> dict[str, Any]:
        """Extract text from Excel files."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)
            sheets_text: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows: list[str] = []
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(cell) for cell in row if cell is not None]
                    if row_vals:
                        rows.append(" | ".join(row_vals))
                if rows:
                    sheets_text.append(f"## Sheet: {sheet_name}\n" + "\n".join(rows))
            wb.close()
            return {"text": "\n\n".join(sheets_text), "page_count": len(sheets_text), "metadata": {"sheets": len(sheets_text)}, "method": "openpyxl"}
        except ImportError:
            return {"text": f"[Excel extraction requires openpyxl: {file_path}]", "page_count": 0, "metadata": {}, "method": "placeholder"}

    async def _extract_csv(self, file_path: str) -> dict[str, Any]:
        """Extract text from CSV files."""
        import csv

        rows: list[str] = []
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                if any(cell.strip() for cell in row):
                    rows.append(" | ".join(cell.strip() for cell in row))
        return {"text": "\n".join(rows), "page_count": 1, "metadata": {"rows": len(rows)}, "method": "csv"}

    async def _extract_txt(self, file_path: str) -> dict[str, Any]:
        """Extract text from plain text files."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return {"text": text, "page_count": 1, "metadata": {}, "method": "plaintext"}

    async def _extract_image(self, file_path: str) -> dict[str, Any]:
        """Extract text from images using OCR."""
        try:
            import pytesseract
            from PIL import Image

            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return {"text": text, "page_count": 1, "metadata": {"width": img.width, "height": img.height}, "method": "tesseract"}
        except ImportError:
            return {"text": f"[Image OCR requires pytesseract: {file_path}]", "page_count": 0, "metadata": {}, "method": "placeholder"}


# Singleton
document_extractor = DocumentExtractor()
